from flask import Flask, render_template, request, send_file
import requests
import re
from bs4 import BeautifulSoup
from os.path import basename
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

app = Flask(__name__)

# Folder untuk menyimpan gambar dan file Word
IMAGE_FOLDER = 'images'
WORD_FOLDER = 'word'
os.makedirs(IMAGE_FOLDER, exist_ok=True)
os.makedirs(WORD_FOLDER, exist_ok=True)

def download_image(img_url):
    img_name = basename(img_url)
    with open(os.path.join(IMAGE_FOLDER, img_name), 'wb') as f:
        image = requests.get(img_url)
        f.write(image.content)
    return img_name 

def get_img_url(soup):
    return soup.find('div', class_='photo__wrap').find('img')['src']

def get_photo_caption(soup):
    return soup.find('div', class_='photo__caption').text.strip()

def get_soup(url):
    page = requests.get(url)
    return BeautifulSoup(page.content, "html.parser")

def get_title(soup):
    return soup.find('h1', class_='read__title').text

def get_clean_content(soup, to_be_erase):
    content = soup.find('div', class_='read__content')
    for data in to_be_erase:
        data.decompose()
    return content

def save_to_word(title, clean_content, filename, caption):
    document = Document()
    judul = document.add_heading(title+'\n', level=1)
    judul.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_picture(os.path.join(IMAGE_FOLDER, filename), width=Inches(6))
    caption_para = document.add_paragraph().add_run(caption+'\n')
    caption_para.font.size = Pt(8)
    caption_para.font.name = 'Arial'
    caption_para.italic = True
    caption_para.font.color.rgb = RGBColor(102, 102, 153)
    isiBerita = document.add_paragraph(clean_content)
    isiBerita.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    file_result = os.path.join(WORD_FOLDER, title + '.docx')
    document.save(file_result)
    return file_result

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        url = request.form['url']
        soup = get_soup(url)    
        title = get_title(soup)
        to_be_erase = soup(['strong', 'i'])
        clean_content = get_clean_content(soup, to_be_erase).text.strip()
        teks_tanpa_baris_kosong = re.sub(r'\n\s*\n+', '\n\n', clean_content)
        img_url = get_img_url(soup)
        filename = download_image(img_url)
        caption = get_photo_caption(soup)
        word_file = save_to_word(title, teks_tanpa_baris_kosong, filename, caption)
        return send_file(word_file, as_attachment=True)
    return render_template('index.html')

if __name__ == '__main__':
    app.run(debug=True)
